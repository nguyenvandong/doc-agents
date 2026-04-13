from __future__ import annotations

import json
from pathlib import Path
import tempfile
import unittest

from docx import Document

from doc_agents.activities import (
    ChunkInput,
    ExtractInput,
    ParseDocumentInput,
    SynthesisInput,
    ValidationInput,
    ValidationReport,
    configure_activity_dependencies,
    extract_business_rules_activity,
    extract_data_schema_activity,
    extract_workflows_activity,
    generate_frontmatter_activity,
    parse_docx_activity,
    persist_markdown_activity,
    render_mermaid_activity,
    semantic_chunk_activity,
    synthesize_markdown_activity,
    validate_markdown_against_chunks_activity,
)
from doc_agents.models import ArtifactRef
from doc_agents.parser import DocxParser


class FakeArtifactRepository:
    def __init__(self) -> None:
        self.calls: list[dict] = []
        self.payloads_by_uri: dict[str, bytes] = {}
        self.latest_versions: dict[tuple[str, str], int] = {}

    def store_bytes(
        self,
        *,
        workflow_id: str,
        document_id: str,
        artifact_type: str,
        payload: bytes,
        content_type: str,
        version: int | None = 1,
    ) -> ArtifactRef:
        resolved_version = version
        if resolved_version is None:
            key = (workflow_id, artifact_type)
            resolved_version = self.latest_versions.get(key, 0) + 1
        self.calls.append(
            {
                "workflow_id": workflow_id,
                "document_id": document_id,
                "artifact_type": artifact_type,
                "payload": payload,
                "content_type": content_type,
                "version": resolved_version,
            }
        )
        artifact = ArtifactRef(
            artifact_id=f"{document_id}-{artifact_type}-v{resolved_version}",
            artifact_type=artifact_type,
            version=resolved_version,
            uri=f"s3://doc-artifacts/{workflow_id}/{artifact_type}/v{resolved_version}/{document_id}-{artifact_type}-v{resolved_version}.bin",
        )
        self.latest_versions[(workflow_id, artifact_type)] = resolved_version
        self.payloads_by_uri[artifact.uri] = payload
        return artifact

    def load_bytes(self, artifact: ArtifactRef) -> bytes:
        return self.payloads_by_uri[artifact.uri]


def build_parsed_document_payload(*texts: str) -> bytes:
    return json.dumps(
        {
            "blocks": [
                {
                    "kind": "paragraph",
                    "text": text,
                    "markdown": text,
                    "style_name": "Normal",
                }
                for text in texts
            ],
            "semantic_html": "<p>semantic</p>",
            "mammoth_messages": [],
        }
    ).encode("utf-8")


def build_chunk_set_payload(*texts: str) -> bytes:
    return json.dumps(
        {
            "chunks": [
                {
                    "chunk_id": f"chunk-{index}",
                    "text": text,
                    "kind": "paragraph",
                    "source_block_indices": [index],
                }
                for index, text in enumerate(texts)
            ]
        }
    ).encode("utf-8")


def build_extraction_payload(*, extraction_kind: str, source_chunk_artifact_id: str, evidence: list[str]) -> bytes:
    return json.dumps(
        {
            "extraction_kind": extraction_kind,
            "source_chunk_artifact_id": source_chunk_artifact_id,
            "source_chunk_count": len(evidence),
            "evidence": evidence,
        }
    ).encode("utf-8")


def persisted_artifact(
    repository: FakeArtifactRepository,
    *,
    document_id: str,
    artifact_type: str,
    payload: bytes,
) -> ArtifactRef:
    return repository.store_bytes(
        workflow_id=document_id,
        document_id=document_id,
        artifact_type=artifact_type,
        payload=payload,
        content_type="application/json",
        version=1,
    )


def build_docx_file(target: Path) -> None:
    document = Document()
    document.add_heading("Eligibility", level=1)
    document.add_paragraph("Applicant must be 18 years old.")
    document.save(target)


class ParseDocxActivityTest(unittest.TestCase):
    def tearDown(self) -> None:
        configure_activity_dependencies(docx_parser=DocxParser(), artifact_repository=None)

    def test_parse_docx_activity_reads_local_file_and_persists_parsed_document(self) -> None:
        repository = FakeArtifactRepository()
        configure_activity_dependencies(docx_parser=DocxParser(), artifact_repository=repository)

        with tempfile.TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / "source.docx"
            build_docx_file(source_path)
            result = parse_docx_activity(
                ParseDocumentInput(
                    document_id="doc-parse",
                    source_uri=str(source_path),
                )
            )

        self.assertEqual(result.artifact_type, "parsed_document")
        self.assertTrue(result.uri.startswith("s3://doc-artifacts/"))
        self.assertEqual(repository.calls[0]["artifact_type"], "parsed_document")
        self.assertIn(b"Eligibility", repository.calls[0]["payload"])

    def test_semantic_chunk_activity_reads_persisted_parsed_document_and_persists_chunk_set(self) -> None:
        repository = FakeArtifactRepository()
        configure_activity_dependencies(docx_parser=DocxParser(), artifact_repository=repository)
        parsed_document = persisted_artifact(
            repository,
            document_id="doc-chunk",
            artifact_type="parsed_document",
            payload=build_parsed_document_payload(
                "Eligibility",
                "Applicant must be 18 years old.",
            ),
        )

        result = semantic_chunk_activity(
            ChunkInput(document_id="doc-chunk", parsed_document=parsed_document)
        )

        self.assertEqual(result.artifact_type, "semantic_chunks")
        self.assertEqual(repository.calls[-1]["artifact_type"], "semantic_chunks")
        persisted_payload = json.loads(repository.calls[-1]["payload"].decode("utf-8"))
        self.assertEqual(
            [chunk["text"] for chunk in persisted_payload["chunks"]],
            ["Eligibility", "Applicant must be 18 years old."],
        )

    def test_extract_activities_read_persisted_chunk_set_and_store_json_artifacts(self) -> None:
        repository = FakeArtifactRepository()
        configure_activity_dependencies(docx_parser=DocxParser(), artifact_repository=repository)
        chunk_set = persisted_artifact(
            repository,
            document_id="doc-extract",
            artifact_type="semantic_chunks",
            payload=build_chunk_set_payload(
                "Field: Customer ID",
                "Rule: Customer ID must be unique.",
                "Workflow: System validates application before approval.",
            ),
        )

        cases = [
            (
                extract_data_schema_activity,
                "data_schema_json",
                {"fields": [{"name": "Customer ID", "source_chunk_id": "chunk-0"}]},
            ),
            (
                extract_business_rules_activity,
                "business_rules_json",
                {"rules": [{"text": "Customer ID must be unique.", "source_chunk_id": "chunk-1"}]},
            ),
            (
                extract_workflows_activity,
                "workflows_json",
                {"steps": [{"text": "System validates application before approval.", "source_chunk_id": "chunk-2"}]},
            ),
        ]
        for activity_fn, expected_artifact_type, expected_payload_slice in cases:
            with self.subTest(activity=expected_artifact_type):
                result = activity_fn(
                    ExtractInput(document_id="doc-extract", chunk_set=chunk_set)
                )

                self.assertEqual(result.artifact_type, expected_artifact_type)
                self.assertEqual(repository.calls[-1]["artifact_type"], expected_artifact_type)
                persisted_payload = json.loads(repository.calls[-1]["payload"].decode("utf-8"))
                self.assertEqual(
                    persisted_payload["source_chunk_artifact_id"],
                    "doc-extract-semantic_chunks-v1",
                )
                for key, value in expected_payload_slice.items():
                    self.assertEqual(persisted_payload[key], value)

    def test_synthesize_markdown_activity_reads_persisted_ir_artifacts_and_persists_markdown(self) -> None:
        repository = FakeArtifactRepository()
        configure_activity_dependencies(docx_parser=DocxParser(), artifact_repository=repository)
        data_schema = persisted_artifact(
            repository,
            document_id="doc-synthesis",
            artifact_type="data_schema_json",
            payload=build_extraction_payload(
                extraction_kind="data_schema",
                source_chunk_artifact_id="doc-synthesis-semantic_chunks-v1",
                evidence=["Customer ID", "Application Date"],
            ),
        )
        business_rules = persisted_artifact(
            repository,
            document_id="doc-synthesis",
            artifact_type="business_rules_json",
            payload=build_extraction_payload(
                extraction_kind="business_rules",
                source_chunk_artifact_id="doc-synthesis-semantic_chunks-v1",
                evidence=["Applicant must be 18 years old."],
            ),
        )
        workflows_json = persisted_artifact(
            repository,
            document_id="doc-synthesis",
            artifact_type="workflows_json",
            payload=build_extraction_payload(
                extraction_kind="workflows",
                source_chunk_artifact_id="doc-synthesis-semantic_chunks-v1",
                evidence=["System validates application before approval."],
            ),
        )

        result = synthesize_markdown_activity(
            SynthesisInput(
                document_id="doc-synthesis",
                data_schema=data_schema,
                business_rules=business_rules,
                workflows=workflows_json,
            )
        )

        self.assertEqual(result.artifact_type, "markdown_draft")
        self.assertEqual(repository.calls[-1]["artifact_type"], "markdown_draft")
        markdown = repository.calls[-1]["payload"].decode("utf-8")
        self.assertIn("# Document Specification", markdown)
        self.assertIn("## Data Schema", markdown)
        self.assertIn("Customer ID", markdown)
        self.assertIn("## Business Rules", markdown)
        self.assertIn("Applicant must be 18 years old.", markdown)

    def test_synthesis_side_artifacts_read_same_persisted_ir_versions(self) -> None:
        repository = FakeArtifactRepository()
        configure_activity_dependencies(docx_parser=DocxParser(), artifact_repository=repository)
        data_schema = persisted_artifact(
            repository,
            document_id="doc-synthesis-side",
            artifact_type="data_schema_json",
            payload=json.dumps(
                {
                    "source_chunk_artifact_id": "doc-synthesis-side-semantic_chunks-v1",
                    "fields": [{"name": "Customer ID", "source_chunk_id": "chunk-0"}],
                    "evidence": ["Customer ID"],
                }
            ).encode("utf-8"),
        )
        business_rules = persisted_artifact(
            repository,
            document_id="doc-synthesis-side",
            artifact_type="business_rules_json",
            payload=json.dumps(
                {
                    "source_chunk_artifact_id": "doc-synthesis-side-semantic_chunks-v1",
                    "rules": [{"text": "Applicant must be 18 years old.", "source_chunk_id": "chunk-1"}],
                    "evidence": ["Applicant must be 18 years old."],
                }
            ).encode("utf-8"),
        )
        workflows_json = persisted_artifact(
            repository,
            document_id="doc-synthesis-side",
            artifact_type="workflows_json",
            payload=json.dumps(
                {
                    "source_chunk_artifact_id": "doc-synthesis-side-semantic_chunks-v1",
                    "steps": [{"text": "System validates application before approval.", "source_chunk_id": "chunk-2"}],
                    "evidence": ["System validates application before approval."],
                }
            ).encode("utf-8"),
        )
        synthesis_input = SynthesisInput(
            document_id="doc-synthesis-side",
            data_schema=data_schema,
            business_rules=business_rules,
            workflows=workflows_json,
        )

        frontmatter = generate_frontmatter_activity(synthesis_input)
        mermaid = render_mermaid_activity(synthesis_input)

        self.assertEqual(frontmatter.version, 1)
        self.assertEqual(mermaid.version, 1)
        self.assertIn("document_id: doc-synthesis-side", repository.load_bytes(frontmatter).decode("utf-8"))
        self.assertIn("flowchart TD", repository.load_bytes(mermaid).decode("utf-8"))

    def test_synthesize_markdown_activity_creates_new_version_on_second_run(self) -> None:
        repository = FakeArtifactRepository()
        configure_activity_dependencies(docx_parser=DocxParser(), artifact_repository=repository)
        data_schema = persisted_artifact(
            repository,
            document_id="doc-synthesis-versioned",
            artifact_type="data_schema_json",
            payload=build_extraction_payload(
                extraction_kind="data_schema",
                source_chunk_artifact_id="doc-synthesis-versioned-semantic_chunks-v1",
                evidence=["Customer ID"],
            ),
        )
        business_rules = persisted_artifact(
            repository,
            document_id="doc-synthesis-versioned",
            artifact_type="business_rules_json",
            payload=build_extraction_payload(
                extraction_kind="business_rules",
                source_chunk_artifact_id="doc-synthesis-versioned-semantic_chunks-v1",
                evidence=["Applicant must be 18 years old."],
            ),
        )
        workflows_json = persisted_artifact(
            repository,
            document_id="doc-synthesis-versioned",
            artifact_type="workflows_json",
            payload=build_extraction_payload(
                extraction_kind="workflows",
                source_chunk_artifact_id="doc-synthesis-versioned-semantic_chunks-v1",
                evidence=["System validates application before approval."],
            ),
        )
        synthesis_input = SynthesisInput(
            document_id="doc-synthesis-versioned",
            data_schema=data_schema,
            business_rules=business_rules,
            workflows=workflows_json,
        )

        first = synthesize_markdown_activity(synthesis_input)
        second = synthesize_markdown_activity(synthesis_input)

        self.assertEqual(first.version, 1)
        self.assertEqual(second.version, 2)
        self.assertNotEqual(first.uri, second.uri)

    def test_persist_markdown_activity_keeps_existing_persisted_artifact(self) -> None:
        repository = FakeArtifactRepository()
        configure_activity_dependencies(docx_parser=DocxParser(), artifact_repository=repository)
        markdown_draft = persisted_artifact(
            repository,
            document_id="doc-markdown",
            artifact_type="markdown_draft",
            payload=b"# Document Specification\n",
        )

        result = persist_markdown_activity(markdown_draft)

        self.assertEqual(result, markdown_draft)

    def test_validate_markdown_activity_reads_persisted_markdown_and_chunks(self) -> None:
        repository = FakeArtifactRepository()
        configure_activity_dependencies(docx_parser=DocxParser(), artifact_repository=repository)
        chunk_set = persisted_artifact(
            repository,
            document_id="doc-validate",
            artifact_type="semantic_chunks",
            payload=build_chunk_set_payload(
                "Field: Customer ID",
                "Rule: Applicant must be 18 years old.",
                "Workflow: System validates application before approval.",
            ),
        )
        markdown_draft = persisted_artifact(
            repository,
            document_id="doc-validate",
            artifact_type="markdown_draft",
            payload=(
                "# Document Specification\n\n"
                "## Data Schema\n\n"
                "- Application Date\n\n"
                "## Business Rules\n\n"
                "- Customer ID must be unique.\n"
            ).encode("utf-8"),
        )

        report = validate_markdown_against_chunks_activity(
            ValidationInput(
                document_id="doc-validate",
                markdown_draft=markdown_draft,
                chunk_set=chunk_set,
            )
        )

        self.assertIsInstance(report, ValidationReport)
        self.assertFalse(report.passed)
        self.assertEqual(
            report.issues,
            [
                "Missing field coverage in markdown: Customer ID",
                "Missing rule coverage in markdown: Applicant must be 18 years old.",
                "Missing workflow coverage in markdown: System validates application before approval.",
            ],
        )


if __name__ == "__main__":
    unittest.main()
