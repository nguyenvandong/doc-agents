from __future__ import annotations

from io import BytesIO
import unittest

from docx import Document

from doc_agents.parser import DocxParser


def build_sample_docx() -> bytes:
    document = Document()
    document.add_heading("Account Opening", level=1)
    document.add_paragraph("Customer submits the form.")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Type"
    table.rows[1].cells[0].text = "customer_id"
    table.rows[1].cells[1].text = "uuid"
    document.add_paragraph("System validates required fields.")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


class DocxParserTest(unittest.TestCase):
    def test_parser_preserves_block_order_and_tables(self) -> None:
        parsed = DocxParser().parse_bytes(build_sample_docx())
        self.assertEqual(
            [block.kind for block in parsed.blocks],
            ["paragraph", "paragraph", "table", "paragraph"],
        )
        self.assertEqual(parsed.blocks[0].style_name, "Heading 1")
        self.assertIn("| Field | Type |", parsed.blocks[2].markdown)
        self.assertEqual(parsed.blocks[3].text, "System validates required fields.")

    def test_parser_generates_semantic_html_with_mammoth(self) -> None:
        parsed = DocxParser().parse_bytes(build_sample_docx())
        self.assertIn("Account Opening", parsed.semantic_html)
        self.assertTrue(parsed.semantic_html.strip().startswith("<h1>"))


if __name__ == "__main__":
    unittest.main()
